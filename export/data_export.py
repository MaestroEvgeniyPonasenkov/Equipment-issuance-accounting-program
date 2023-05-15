from docx import Document
import openpyxl


def export_to_xlsx(data: list[dict], filename: str) -> None:
    """
    Saves the given data as xlsx file.

    Args:
    - data (list): A list of dictionaries representing the data to be exported.
    - filename (str): The filename (without extension) for the output file.

    Returns:
    - None.
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    headers = list(data[0].keys())
    for col_num, header in enumerate(headers, 1):
        sheet.cell(row=1, column=col_num, value=header)
    for row_num, row_data in enumerate(data, 2):
        for col_num, col_name in enumerate(headers, 1):
            sheet.cell(row=row_num, column=col_num, value=row_data[col_name])
    workbook.save(f"{filename}.xlsx")
    print("Экспорт данных успешно выполнен!")


def export_to_docx(data: list[dict], filename: str) -> None:
    """
    Saves the given data as a docx file.

    Args:
        data (list): A list of dictionaries representing the data to be exported.
        filename (str): The filename (without extension) for the output file.

    Returns:
        None
    """
    document = Document()
    for i, row_data in enumerate(data, 1):
        paragraph = document.add_paragraph()
        paragraph.add_run(f'Данные №{i}:\n').bold = True
        for key, value in row_data.items():
            paragraph.add_run(f'{key}: {value}\n').bold = False
        paragraph.add_run('\n')
    document.save(f"{filename}.docx")
    print("Экспорт данных успешно выполнен!")
