import csv
import docx
from docx.shared import Pt


def xlsx_xprt(name: str, header: list, data: list) -> None:
    """
    Saves the given data from the database as an excel table.

    Args:
    - name (str): Name of the output file.
    - header (list): A list of strings representing the header for the table.
    - data (list): A list of lists representing the data for the table.

    Returns:
    - None.
    """
    with open(f"{name}.csv", "w", newline='') as xlsx_file:
        writer = csv.writer(xlsx_file)
        writer.writerow(header)
        writer.writerows(data)


def docx_xprt(name: str, header: list, data: list) -> None:
    """
    Saves the given data from the database as a docx document.

    Args:
    - name (str): Name of the output file.
    - header (list): A list of strings representing the header for the document.
    - data (list): A list of lists representing the data for the document.

    Returns:
    - None.
    """
    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(14)
    for h in header:
        doc.add_paragraph(h)
    for d in data:
        if isinstance(d, list) or isinstance(d, tuple):
            d = ", ".join(str(i) for i in d)
        doc.add_paragraph(d)
    doc.save(f"{name}.docx")
